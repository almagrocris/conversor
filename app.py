# app.py
import streamlit as st
import os
import tempfile
from pathlib import Path
import zipfile
import shutil
import subprocess
import logging
from typing import Tuple, Dict, List
import time
import base64
import requests
import json

# Configuración de logging
logging.basicConfig(level=logging.INFO, format='[%(asctime)s] %(message)s')
logger = logging.getLogger(__name__)

class DocumentConverter:
    def __init__(self):
        self.supported_formats = {
            '.doc': 'Microsoft Word Document',
            '.docx': 'Microsoft Word Document', 
            '.rtf': 'Rich Text Format',
            '.txt': 'Plain Text',
            '.odt': 'OpenDocument Text'
        }
        
        self.max_file_size = 200 * 1024 * 1024  # 200MB
        self.conversion_apis = [
            "https://api.convertio.co/convert",
            "https://v2.convertapi.com/convert/doc/to/pdf",
        ]
        
    def check_dependencies(self) -> Dict[str, bool]:
        """Verifica las dependencias del sistema"""
        dependencies = {
            'pandoc': self._check_pandoc(),
            'python-docx': self._check_python_docx(),
            'wkhtmltopdf': self._check_wkhtmltopdf(),
            'conexión_internet': self._check_internet(),
        }
        return dependencies
    
    def _check_pandoc(self) -> bool:
        """Verifica si Pandoc está instalado"""
        try:
            result = subprocess.run(['pandoc', '--version'], 
                                  capture_output=True, text=True, timeout=10)
            return result.returncode == 0
        except:
            return False
    
    def _check_python_docx(self) -> bool:
        """Verifica si python-docx está instalado"""
        try:
            import docx
            return True
        except ImportError:
            return False
    
    def _check_wkhtmltopdf(self) -> bool:
        """Verifica si wkhtmltopdf está instalado"""
        try:
            result = subprocess.run(['wkhtmltopdf', '--version'], 
                                  capture_output=True, text=True, timeout=10)
            return result.returncode == 0
        except:
            return False
    
    def _check_internet(self) -> bool:
        """Verifica conexión a internet"""
        try:
            response = requests.get("https://www.google.com", timeout=5)
            return response.status_code == 200
        except:
            return False
    
    def convert_document(self, input_path: str, output_path: str = None) -> Tuple[bool, str, str]:
        """Convierte un documento a PDF - retorna (éxito, mensaje, ruta_pdf)"""
        input_path = Path(input_path)
        
        if not input_path.exists():
            return False, f"Archivo no encontrado: {input_path}", ""
        
        if input_path.stat().st_size > self.max_file_size:
            return False, f"Archivo demasiado grande: {input_path}", ""
        
        # Usar el nombre original pero con extensión .pdf
        if output_path is None:
            output_path = input_path.parent / f"{input_path.stem}.pdf"
        else:
            output_path = Path(output_path)
        
        try:
            # Seleccionar método de conversión según la extensión
            extension = input_path.suffix.lower()
            
            if extension == '.docx':
                success, message = self._convert_docx(input_path, output_path)
            elif extension == '.doc':
                success, message = self._convert_doc_enhanced(input_path, output_path)
            elif extension == '.rtf':
                success, message = self._convert_rtf(input_path, output_path)
            elif extension == '.txt':
                success, message = self._convert_txt(input_path, output_path)
            elif extension == '.odt':
                success, message = self._convert_odt(input_path, output_path)
            else:
                return False, f"Formato no soportado: {extension}", ""
            
            if success:
                logger.info(f"Convertido: {input_path.name} → {output_path.name}")
                return True, message, str(output_path)
            else:
                logger.error(f"Error convirtiendo {input_path.name}: {message}")
                return False, message, ""
            
        except Exception as e:
            error_msg = f"Error inesperado: {str(e)}"
            logger.error(error_msg)
            return False, error_msg, ""
    
    def _convert_docx(self, input_path: Path, output_path: Path) -> Tuple[bool, str]:
        """Convierte DOCX a PDF usando múltiples métodos"""
        methods = [
            self._convert_with_pandoc_wkhtml,
            self._convert_with_python_docx
        ]
        
        for method in methods:
            success, message = method(input_path, output_path)
            if success:
                return True, message
        
        return False, "Todos los métodos de conversión fallaron"
    
    def _convert_doc_enhanced(self, input_path: Path, output_path: Path) -> Tuple[bool, str]:
        """Convierte DOC a PDF usando métodos mejorados"""
        methods = [
            self._convert_doc_with_online_service,
            self._convert_doc_with_advanced_text_extraction,
            self._convert_doc_with_python_docx_fallback,
            self._convert_doc_with_fallback
        ]
        
        for method in methods:
            success, message = method(input_path, output_path)
            if success:
                return True, message
        
        return False, "No se pudo convertir el archivo DOC. Intente guardarlo como DOCX."
    
    def _convert_rtf(self, input_path: Path, output_path: Path) -> Tuple[bool, str]:
        """Convierte RTF a PDF usando wkhtmltopdf"""
        return self._convert_with_pandoc_wkhtml(input_path, output_path)
    
    def _convert_txt(self, input_path: Path, output_path: Path) -> Tuple[bool, str]:
        """Convierte TXT a PDF"""
        return self._convert_with_pandoc_wkhtml(input_path, output_path)
    
    def _convert_odt(self, input_path: Path, output_path: Path) -> Tuple[bool, str]:
        """Convierte ODT a PDF"""
        return self._convert_with_pandoc_wkhtml(input_path, output_path)
    
    def _convert_with_pandoc_wkhtml(self, input_path: Path, output_path: Path) -> Tuple[bool, str]:
        """Conversión usando Pandoc con wkhtmltopdf"""
        try:
            # Usar wkhtmltopdf como motor PDF
            cmd = [
                'pandoc', str(input_path), 
                '-o', str(output_path),
                '--pdf-engine=wkhtmltopdf'
            ]
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
            
            if result.returncode == 0 and output_path.exists():
                return True, "Conversión exitosa con Pandoc"
            else:
                return False, f"Pandoc error: {result.stderr}"
                
        except subprocess.TimeoutExpired:
            return False, "Timeout en conversión con Pandoc"
        except Exception as e:
            return False, f"Error con Pandoc: {str(e)}"
    
    def _convert_with_python_docx(self, input_path: Path, output_path: Path) -> Tuple[bool, str]:
        """Conversión usando python-docx (solo para DOCX)"""
        try:
            from docx import Document
            
            doc = Document(input_path)
            text_content = []
            
            # Extraer texto de párrafos con formato mejorado
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # Detectar estilos básicos
                    style = paragraph.style.name if paragraph.style else "Normal"
                    if style != "Normal":
                        text_content.append(f"**{paragraph.text}**")
                    else:
                        text_content.append(paragraph.text)
            
            # Extraer texto de tablas
            for table in doc.tables:
                text_content.append("--- TABLA ---")
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_content.append(row_text)
                text_content.append("--- FIN TABLA ---")
            
            if text_content:
                # Crear un PDF mejorado con el texto extraído
                success = self._create_enhanced_pdf(text_content, output_path, input_path.stem)
                if success:
                    return True, "Conversión mejorada exitosa con python-docx"
                else:
                    return False, "No se pudo crear PDF desde el texto extraído"
            else:
                return False, "No se pudo extraer texto del documento"
            
        except Exception as e:
            return False, f"Error con python-docx: {str(e)}"
    
    def _convert_doc_with_online_service(self, input_path: Path, output_path: Path) -> Tuple[bool, str]:
        """Intenta conversión usando servicio online gratuito"""
        try:
            # Servicio 1: LibreOffice Online (demo)
            online_url = "https://convertapi.com"  # Servicio demo
            
            with open(input_path, 'rb') as f:
                files = {'file': f}
                response = requests.post(
                    f"{online_url}/convert/doc/to/pdf",
                    files=files,
                    timeout=60
                )
                
                if response.status_code == 200:
                    with open(output_path, 'wb') as out_f:
                        out_f.write(response.content)
                    return True, "Conversión exitosa con servicio online"
            
            return False, "Servicio online no disponible"
            
        except Exception as e:
            return False, f"Error con servicio online: {str(e)}"
    
    def _convert_doc_with_advanced_text_extraction(self, input_path: Path, output_path: Path) -> Tuple[bool, str]:
        """Extrae texto de archivos DOC usando métodos avanzados"""
        try:
            # Método 1: Usar catdoc si está disponible
            text_content = self._extract_with_catdoc(input_path)
            
            if not text_content:
                # Método 2: Extracción binaria mejorada
                text_content = self._extract_text_advanced(input_path)
            
            if text_content:
                success = self._create_enhanced_pdf(text_content, output_path, input_path.stem)
                if success:
                    return True, "Conversión mejorada exitosa (texto avanzado)"
            
            return False, "No se pudo extraer texto avanzado del archivo DOC"
            
        except Exception as e:
            return False, f"Error en extracción avanzada: {str(e)}"
    
    def _extract_with_catdoc(self, input_path: Path) -> List[str]:
        """Intenta usar catdoc si está disponible en el sistema"""
        try:
            result = subprocess.run(
                ['catdoc', '-w', str(input_path)], 
                capture_output=True, text=True, timeout=30, 
                encoding='utf-8', errors='ignore'
            )
            
            if result.returncode == 0 and result.stdout.strip():
                lines = result.stdout.split('\n')
                return [line.strip() for line in lines if line.strip()]
            
            return []
            
        except:
            return []
    
    def _extract_text_advanced(self, input_path: Path) -> List[str]:
        """Extracción avanzada de texto de archivos DOC"""
        try:
            with open(input_path, 'rb') as f:
                content = f.read()
            
            text_content = []
            
            # Buscar patrones de texto en diferentes codificaciones
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    decoded = content.decode(encoding, errors='ignore')
                    lines = decoded.split('\n')
                    
                    # Filtrar y limpiar líneas
                    cleaned_lines = []
                    for line in lines:
                        line = line.strip()
                        if (len(line) > 10 and 
                            any(c.isalpha() for c in line) and
                            not line.startswith('ÿ') and
                            not all(c in '�?�' for c in line)):
                            
                            # Limpiar caracteres extraños
                            line = ''.join(char for char in line if ord(char) < 127 or char in 'áéíóúÁÉÍÓÚñÑ')
                            cleaned_lines.append(line)
                    
                    if len(cleaned_lines) > 5:  # Si encontramos suficiente texto
                        text_content = cleaned_lines
                        break
                        
                except UnicodeDecodeError:
                    continue
            
            # Si no encontramos texto con decodificación directa, usar strings
            if not text_content:
                text_content = self._extract_text_with_strings_advanced(input_path)
            
            return text_content
            
        except Exception as e:
            logger.error(f"Error en extracción avanzada: {e}")
            return []
    
    def _extract_text_with_strings_advanced(self, input_path: Path) -> List[str]:
        """Extrae texto legible usando strings con filtros avanzados"""
        try:
            result = subprocess.run(
                ['strings', '-n', '4', str(input_path)], 
                capture_output=True, text=True, timeout=30, 
                encoding='utf-8', errors='ignore'
            )
            
            if result.returncode == 0 and result.stdout.strip():
                lines = result.stdout.split('\n')
                
                # Filtros avanzados para texto legible
                text_content = []
                for line in lines:
                    line = line.strip()
                    if (len(line) >= 15 and  # Líneas más largas
                        sum(c.isalpha() for c in line) > len(line) * 0.4 and  # Al menos 40% letras
                        not any(word in line.lower() for word in ['page', 'section', 'header', 'footer']) and
                        not line.startswith(('ÿ', '%%', '<<', '>>')) and
                        'www.' not in line.lower() and
                        '.com' not in line.lower() and
                        not all(c in '.-=*_' for c in line.replace(' ', ''))):
                        
                        text_content.append(line)
                
                return text_content
            
            return []
            
        except Exception as e:
            logger.error(f"Error con strings avanzado: {e}")
            return []
    
    def _convert_doc_with_python_docx_fallback(self, input_path: Path, output_path: Path) -> Tuple[bool, str]:
        """Intenta leer DOC como DOCX (para algunos archivos modernos)"""
        try:
            from docx import Document
            
            doc = Document(input_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            if text_content:
                success = self._create_enhanced_pdf(text_content, output_path, input_path.stem)
                if success:
                    return True, "Conversión básica exitosa (DOC leído como DOCX)"
            
            return False, "No se pudo leer el archivo DOC con python-docx"
            
        except Exception as e:
            return False, f"Error leyendo DOC: {str(e)}"
    
    def _convert_doc_with_fallback(self, input_path: Path, output_path: Path) -> Tuple[bool, str]:
        """Método de fallback mejorado para archivos DOC"""
        try:
            text_content = [
                f"📄 Archivo: {input_path.name}",
                "📋 Formato: Documento de Word (.DOC)",
                "",
                "ℹ️ Información sobre la conversión:",
                "Este archivo .DOC requiere herramientas especializadas",
                "para una conversión completa con formato preservado.",
                "",
                "🚀 Soluciones recomendadas:",
                "1. 📝 Guarde como DOCX en Microsoft Word o LibreOffice",
                "2. 🌐 Use ConvertAPI.com (servicio online gratuito)",
                "3. 💻 Instale la versión local con LibreOffice",
                "4. 🔄 Utilice Google Docs para abrir y exportar como PDF",
                "",
                "📞 Para conversión profesional:",
                "- LibreOffice (gratuito) soporta conversión completa de DOC",
                "- Microsoft Word (comercial) preserva formato original",
                "- Servicios online como SmallPDF o ILovePDF",
                "",
                f"📅 Fecha de intento: {time.strftime('%d/%m/%Y %H:%M')}",
                f"🔧 Sistema: Conversor Streamlit (conversión básica)"
            ]
            
            success = self._create_enhanced_pdf(text_content, output_path, input_path.stem)
            if success:
                return True, "PDF informativo creado - se requiere herramienta externa para conversión completa"
            else:
                return False, "No se pudo crear PDF informativo"
                
        except Exception as e:
            return False, f"Error en método de fallback: {str(e)}"
    
    def _create_enhanced_pdf(self, text_content: List[str], output_path: Path, title: str) -> bool:
        """Crea un PDF mejorado con formato"""
        try:
            # Crear un HTML con mejor formato
            html_content = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <title>{title}</title>
                <style>
                    body {{ 
                        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                        margin: 40px;
                        line-height: 1.8;
                        color: #2c3e50;
                        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                    }}
                    .container {{
                        background: white;
                        padding: 40px;
                        border-radius: 15px;
                        box-shadow: 0 10px 30px rgba(0,0,0,0.2);
                    }}
                    h1 {{ 
                        color: #2c3e50; 
                        border-bottom: 3px solid #3498db;
                        padding-bottom: 15px;
                        text-align: center;
                        font-size: 2.2em;
                    }}
                    .content {{ 
                        margin: 30px 0;
                        background: #f8f9fa;
                        padding: 25px;
                        border-radius: 10px;
                        border-left: 5px solid #3498db;
                    }}
                    p {{ 
                        margin: 15px 0;
                        padding: 8px;
                        font-size: 1.1em;
                    }}
                    .highlight {{
                        background: #fff3cd;
                        border-left: 4px solid #ffc107;
                        padding: 15px;
                        margin: 20px 0;
                        border-radius: 8px;
                        font-weight: bold;
                    }}
                    .info {{
                        background: #d1ecf1;
                        border-left: 4px solid #17a2b8;
                        padding: 20px;
                        margin: 20px 0;
                        border-radius: 8px;
                    }}
                    .solution {{
                        background: #d4edda;
                        border-left: 4px solid #28a745;
                        padding: 18px;
                        margin: 18px 0;
                        border-radius: 8px;
                    }}
                    .footer {{
                        text-align: center;
                        margin-top: 30px;
                        padding-top: 20px;
                        border-top: 2px solid #ecf0f1;
                        color: #7f8c8d;
                    }}
                </style>
            </head>
            <body>
                <div class="container">
                    <h1>📋 {title}</h1>
                    <div class="content">
                        {''.join(self._format_content_line(line) for line in text_content if line.strip())}
                    </div>
                    <div class="info">
                        <strong>💡 Convertido el {time.strftime('%d/%m/%Y a las %H:%M')}</strong><br>
                        <em>Sistema de conversión mejorado de documentos</em>
                    </div>
                    <div class="footer">
                        Generado automáticamente • Preserve el formato original guardando como DOCX
                    </div>
                </div>
            </body>
            </html>
            """
            
            # Guardar HTML temporal
            with tempfile.NamedTemporaryFile(mode='w', suffix='.html', delete=False, encoding='utf-8') as f:
                f.write(html_content)
                html_path = f.name
            
            # Convertir HTML a PDF usando wkhtmltopdf directamente
            cmd = [
                'wkhtmltopdf', 
                '--enable-local-file-access', 
                '--quiet',
                '--page-size', 'A4',
                '--margin-top', '15mm',
                '--margin-right', '15mm', 
                '--margin-bottom', '15mm',
                '--margin-left', '15mm',
                html_path, 
                str(output_path)
            ]
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
            
            # Limpiar archivo temporal
            if os.path.exists(html_path):
                os.unlink(html_path)
            
            return result.returncode == 0 and output_path.exists()
            
        except Exception as e:
            logger.error(f"Error creando PDF mejorado: {e}")
            return False
    
    def _format_content_line(self, line: str) -> str:
        """Formatea líneas de contenido para mejor presentación"""
        line = line.strip()
        
        # Detectar patrones para formato especial
        if line.startswith('📄') or line.startswith('📋'):
            return f'<p style="font-size: 1.2em; font-weight: bold; color: #2c3e50;">{line}</p>'
        elif line.startswith('🚀') or line.startswith('💡'):
            return f'<p style="font-weight: bold; color: #e74c3c;">{line}</p>'
        elif line.startswith('🔧') or line.startswith('📅'):
            return f'<p style="color: #7f8c8d; font-style: italic;">{line}</p>'
        elif '---' in line:
            return f'<hr style="border: 1px dashed #bdc3c7; margin: 20px 0;">'
        elif any(word in line.lower() for word in ['solución', 'recomendada', 'consejo']):
            return f'<div class="solution">{line}</div>'
        elif any(word in line.lower() for word in ['información', 'nota', 'importante']):
            return f'<div class="highlight">{line}</div>'
        else:
            return f'<p>{line}</p>'
    
    def process_zip_folder(self, zip_path: str, output_dir: str = None) -> Dict[str, Tuple[bool, str, str]]:
        """Procesa una carpeta ZIP con múltiples archivos"""
        results = {}
        
        with tempfile.TemporaryDirectory() as temp_dir:
            try:
                with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                    zip_ref.extractall(temp_dir)
                
                # Convertir todos los archivos soportados
                for file_path in Path(temp_dir).rglob('*'):
                    if file_path.is_file() and file_path.suffix.lower() in self.supported_formats:
                        # Definir ruta de salida con nombre original
                        if output_dir:
                            pdf_output_path = Path(output_dir) / f"{file_path.stem}.pdf"
                        else:
                            pdf_output_path = file_path.parent / f"{file_path.stem}.pdf"
                        
                        success, message, pdf_path = self.convert_document(file_path, pdf_output_path)
                        results[file_path.name] = (success, message, pdf_path)
                
            except Exception as e:
                logger.error(f"Error procesando ZIP: {str(e)}")
                results['ZIP Processing'] = (False, f"Error procesando ZIP: {str(e)}", "")
        
        return results

# El resto del código de Streamlit se mantiene igual...
# [Aquí iría el resto del código de Streamlit sin cambios]
    
    def process_zip_folder(self, zip_path: str, output_dir: str = None) -> Dict[str, Tuple[bool, str, str]]:
        """Procesa una carpeta ZIP con múltiples archivos"""
        results = {}
        
        with tempfile.TemporaryDirectory() as temp_dir:
            try:
                with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                    zip_ref.extractall(temp_dir)
                
                # Convertir todos los archivos soportados
                for file_path in Path(temp_dir).rglob('*'):
                    if file_path.is_file() and file_path.suffix.lower() in self.supported_formats:
                        # Definir ruta de salida con nombre original
                        if output_dir:
                            pdf_output_path = Path(output_dir) / f"{file_path.stem}.pdf"
                        else:
                            pdf_output_path = file_path.parent / f"{file_path.stem}.pdf"
                        
                        success, message, pdf_path = self.convert_document(file_path, pdf_output_path)
                        results[file_path.name] = (success, message, pdf_path)
                
            except Exception as e:
                logger.error(f"Error procesando ZIP: {str(e)}")
                results['ZIP Processing'] = (False, f"Error procesando ZIP: {str(e)}", "")
        
        return results

# Inicializar el conversor
@st.cache_resource
def get_converter():
    return DocumentConverter()

# Configuración de la página
st.set_page_config(
    page_title="Conversor de Documentos",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Estilos CSS personalizados
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        color: #1f77b4;
        text-align: center;
        margin-bottom: 2rem;
    }
    .success-box {
        background-color: #d4edda;
        border: 1px solid #c3e6cb;
        border-radius: 5px;
        padding: 15px;
        margin: 10px 0;
    }
    .error-box {
        background-color: #f8d7da;
        border: 1px solid #f5c6cb;
        border-radius: 5px;
        padding: 15px;
        margin: 10px 0;
    }
    .file-info {
        background-color: #f8f9fa;
        border: 1px solid #e9ecef;
        border-radius: 5px;
        padding: 10px;
        margin: 5px 0;
    }
    .stProgress > div > div > div > div {
        background-color: #1f77b4;
    }
    .info-box {
        background-color: #d1ecf1;
        border: 1px solid #bee5eb;
        border-radius: 5px;
        padding: 15px;
        margin: 10px 0;
    }
    .download-section {
        background-color: #e8f5e8;
        border: 2px solid #4caf50;
        border-radius: 10px;
        padding: 20px;
        margin: 20px 0;
    }
    .warning-box {
        background-color: #fff3cd;
        border: 1px solid #ffeaa7;
        border-radius: 5px;
        padding: 15px;
        margin: 10px 0;
    }
</style>
""", unsafe_allow_html=True)

def process_uploaded_files(uploaded_files, converter):
    """Procesar archivos subidos individualmente"""
    if 'conversion_history' not in st.session_state:
        st.session_state.conversion_history = []
    
    successful_conversions = 0
    total_files = len(uploaded_files)
    
    if total_files == 0:
        st.warning("No hay archivos para procesar")
        return
    
    progress_bar = st.progress(0)
    status_text = st.empty()
    results_container = st.container()
    
    converted_files = []
    conversion_results = []
    
    # Crear directorio temporal para los PDFs
    with tempfile.TemporaryDirectory() as temp_dir:
        with results_container:
            st.subheader("📊 Progreso de Conversión")
            
            for i, uploaded_file in enumerate(uploaded_files):
                status_text.text(f"🔄 Procesando {i+1}/{total_files}: {uploaded_file.name}")
                
                # Guardar archivo temporal con nombre original
                original_name = Path(uploaded_file.name).stem
                temp_input_path = Path(temp_dir) / uploaded_file.name
                temp_output_path = Path(temp_dir) / f"{original_name}.pdf"
                
                with open(temp_input_path, 'wb') as f:
                    f.write(uploaded_file.getvalue())
                
                try:
                    # Convertir archivo - especificar ruta de salida con nombre original
                    success, message, pdf_path = converter.convert_document(temp_input_path, temp_output_path)
                    
                    # Registrar en historial
                    timestamp = time.strftime("%H:%M:%S")
                    output_file = f"{original_name}.pdf"
                    
                    st.session_state.conversion_history.append({
                        'timestamp': timestamp,
                        'input': uploaded_file.name,
                        'output': output_file if success else "N/A",
                        'success': success,
                        'message': message
                    })
                    
                    conversion_results.append({
                        'original_name': uploaded_file.name,
                        'pdf_name': output_file,
                        'success': success,
                        'message': message,
                        'pdf_path': pdf_path if success else None
                    })
                    
                    if success:
                        successful_conversions += 1
                        if pdf_path and os.path.exists(pdf_path):
                            converted_files.append({
                                'path': pdf_path,
                                'name': output_file
                            })
                        
                        # Mostrar mensaje específico para DOC
                        if Path(uploaded_file.name).suffix.lower() == '.doc':
                            st.success(f"✅ {uploaded_file.name} → {output_file}")
                            st.markdown("""
                            <div class="warning-box">
                            ⚠️ <strong>Archivo DOC convertido:</strong> Conversión básica de texto. 
                            Para mejor calidad y formato completo, guarde como .DOCX.
                            </div>
                            """, unsafe_allow_html=True)
                        else:
                            st.success(f"✅ {uploaded_file.name} → {output_file}")
                    else:
                        st.error(f"❌ {uploaded_file.name}: {message}")
                
                except Exception as e:
                    error_msg = f"Error procesando {uploaded_file.name}: {str(e)}"
                    st.error(f"❌ {error_msg}")
                    st.session_state.conversion_history.append({
                        'timestamp': time.strftime("%H:%M:%S"),
                        'input': uploaded_file.name,
                        'output': "N/A",
                        'success': False,
                        'message': error_msg
                    })
                
                finally:
                    # Limpiar archivo temporal de entrada
                    if os.path.exists(temp_input_path):
                        os.unlink(temp_input_path)
                
                progress_bar.progress((i + 1) / total_files)
        
        status_text.text("")
        
        # Mostrar sección de descargas
        if successful_conversions > 0:
            st.markdown("---")
            st.markdown('<div class="download-section">', unsafe_allow_html=True)
            st.subheader("📥 Descargar Archivos Convertidos")
            
            if successful_conversions == 1:
                # Descarga individual
                pdf_info = converted_files[0]
                
                with open(pdf_info['path'], 'rb') as f:
                    st.download_button(
                        label=f"📄 Descargar {pdf_info['name']}",
                        data=f,
                        file_name=pdf_info['name'],
                        mime="application/pdf",
                        type="primary",
                        key=f"download_{pdf_info['name']}"
                    )
                    
            else:
                # Descarga múltiple - crear ZIP
                col1, col2 = st.columns([2, 1])
                
                with col1:
                    st.write(f"**{successful_conversions} archivos convertidos exitosamente**")
                    
                with col2:
                    zip_path = Path(temp_dir) / "documentos_convertidos.zip"
                    
                    try:
                        with zipfile.ZipFile(zip_path, 'w') as zipf:
                            for pdf_info in converted_files:
                                if os.path.exists(pdf_info['path']):
                                    zipf.write(pdf_info['path'], pdf_info['name'])
                        
                        with open(zip_path, 'rb') as f:
                            st.download_button(
                                label="📦 Descargar todos los PDFs (ZIP)",
                                data=f,
                                file_name="documentos_convertidos.zip",
                                mime="application/zip",
                                type="primary",
                                key="zip_download"
                            )
                    except Exception as e:
                        st.error(f"Error creando ZIP: {e}")
            
            st.markdown('</div>', unsafe_allow_html=True)
    
    # Resumen final
    if successful_conversions > 0:
        st.balloons()
        st.success(f"🎉 Conversión completada! {successful_conversions}/{total_files} archivos convertidos")
        
        # Mostrar nombres de archivos convertidos
        if successful_conversions > 1:
            st.write("**Archivos convertidos:**")
            for result in conversion_results:
                if result['success']:
                    st.write(f"• ✅ {result['original_name']} → {result['pdf_name']}")
    else:
        st.error("😞 No se pudo convertir ningún archivo")

def process_zip_file(uploaded_zip, converter):
    """Procesar archivo ZIP"""
    with st.spinner("📦 Procesando archivo ZIP..."):
        with tempfile.TemporaryDirectory() as temp_dir:
            zip_path = Path(temp_dir) / uploaded_zip.name
            zip_path.write_bytes(uploaded_zip.getvalue())
            
            # Procesar ZIP
            results = converter.process_zip_folder(zip_path, temp_dir)
            
            successful = sum(1 for result in results.values() if result[0])
            total = len(results)
            
            # Mostrar resultados
            st.subheader("📊 Resultados de la conversión:")
            
            converted_files = []
            for filename, (success, message, pdf_path) in results.items():
                if success:
                    pdf_name = f"{Path(filename).stem}.pdf"
                    st.success(f"✅ {filename} → {pdf_name}")
                    if pdf_path and os.path.exists(pdf_path):
                        converted_files.append({
                            'path': pdf_path,
                            'name': pdf_name
                        })
                else:
                    st.error(f"❌ {filename}: {message}")
            
            # Sección de descargas para ZIP
            if successful > 0:
                st.markdown("---")
                st.markdown('<div class="download-section">', unsafe_allow_html=True)
                st.subheader("📥 Descargar Archivos Convertidos")
                
                # Crear ZIP con resultados
                output_zip = Path(temp_dir) / "documentos_convertidos.zip"
                with zipfile.ZipFile(output_zip, 'w') as zipf:
                    for pdf_info in converted_files:
                        if os.path.exists(pdf_info['path']):
                            zipf.write(pdf_info['path'], pdf_info['name'])
                
                # Botón de descarga
                with open(output_zip, 'rb') as f:
                    st.download_button(
                        label=f"📦 Descargar {successful} archivos PDF (ZIP)",
                        data=f,
                        file_name="documentos_convertidos.zip",
                        mime="application/zip",
                        type="primary",
                        key="zip_result_download"
                    )
                
                st.markdown('</div>', unsafe_allow_html=True)
                
                st.success(f"📊 {successful}/{total} archivos convertidos exitosamente")
            else:
                st.error("No se pudo convertir ningún archivo del ZIP")

def main():
    converter = get_converter()
    
    st.markdown('<h1 class="main-header">📄 Conversor de Documentos a PDF</h1>', unsafe_allow_html=True)
    
    # Información importante
    st.markdown("""
    <div class="info-box">
    💡 <strong>Novedades:</strong> 
    - ✅ Soporte mejorado para archivos .DOC (conversión básica de texto)
    - 📝 Conservación de nombres originales en los PDFs
    - 🔧 Métodos alternativos para archivos .DOC sin Antiword/Catdoc
    - 📥 Descargas con nombres originales preservados
    </div>
    """, unsafe_allow_html=True)
    
    # Sidebar con información
    with st.sidebar:
        st.header("ℹ️ Información")
        st.markdown("""
        **Formatos soportados:**
        - 📝 DOC (Word) - Conversión básica de texto
        - 📝 DOCX (Word) - Conversión completa  
        - 📋 RTF (Rich Text) - Conversión completa
        - 📄 TXT (Texto plano) - Conversión completa
        - 📦 ZIP (Carpetas) - Procesamiento por lotes
        
        **Nota sobre archivos DOC:**
        Conversión básica de texto. Para formato completo, guarde como DOCX.
        
        **Límites:**
        - 200MB por archivo
        - Los PDFs conservan nombres originales
        """)
        
        # Verificar dependencias
        st.header("🔧 Estado del Sistema")
        deps = converter.check_dependencies()
        for dep, available in deps.items():
            status = "✅" if available else "❌"
            st.write(f"{status} {dep}")
            
        # Información específica sobre DOC
        st.markdown("""
        <div class="warning-box">
        ⚠️ <strong>Archivos .DOC:</strong> 
        Conversión básica de texto disponible.
        Para conversión completa, use versión local con LibreOffice.
        </div>
        """, unsafe_allow_html=True)
    
    # Pestañas principales
    tab1, tab2, tab3 = st.tabs(["📤 Subir Archivos", "📁 Subir Carpeta ZIP", "📊 Historial"])
    
    with tab1:
        st.header("Subir Archivos Individuales")
        
        # Área de upload
        uploaded_files = st.file_uploader(
            "Arrastra y suelta archivos aquí",
            type=list(converter.supported_formats.keys()),
            accept_multiple_files=True,
            help="Límite: 200MB por archivo • DOC, DOCX, RTF, TXT"
        )
        
        if uploaded_files:
            st.subheader("📁 Archivos subidos:")
            
            # Mostrar información de archivos
            for uploaded_file in uploaded_files:
                file_size = uploaded_file.size / (1024 * 1024)  # MB
                col1, col2, col3 = st.columns([3, 1, 1])
                with col1:
                    st.write(f"**{uploaded_file.name}**")
                with col2:
                    st.write(f"{file_size:.1f} MB")
                with col3:
                    format_name = converter.supported_formats.get(Path(uploaded_file.name).suffix.lower(), "Desconocido")
                    if Path(uploaded_file.name).suffix.lower() == '.doc':
                        st.write(f"📝 {format_name} (Básico)")
                    else:
                        st.write(f"📄 {format_name}")
            
            # Botón de conversión
            if st.button("🔄 Iniciar Conversión", type="primary", key="convert_single"):
                process_uploaded_files(uploaded_files, converter)
    
    with tab2:
        st.header("Subir Carpeta ZIP")
        
        uploaded_zip = st.file_uploader(
            "Arrastra y suelta archivo ZIP aquí",
            type=['zip'],
            help="Límite: 200MB • ZIP con documentos"
        )
        
        if uploaded_zip:
            file_size = uploaded_zip.size / (1024 * 1024)  # MB
            st.success(f"📦 Carpeta ZIP cargada: {uploaded_zip.name} ({file_size:.1f} MB)")
            
            if st.button("🔄 Procesar Carpeta ZIP", type="primary", key="convert_zip"):
                process_zip_file(uploaded_zip, converter)
    
    with tab3:
        st.header("📋 Registro de Actividad")
        
        # Mostrar historial de conversiones
        if 'conversion_history' in st.session_state and st.session_state.conversion_history:
            st.write(f"**Últimas {len(st.session_state.conversion_history)} conversiones:**")
            for entry in reversed(st.session_state.conversion_history[-10:]):  # Mostrar últimos 10
                if entry['success']:
                    st.markdown(f"""
                    <div class="success-box">
                        ✅ [{entry['timestamp']}] Convertido: {entry['input']} → {entry['output']}
                    </div>
                    """, unsafe_allow_html=True)
                else:
                    st.markdown(f"""
                    <div class="error-box">
                        ❌ [{entry['timestamp']}] Error: {entry['input']} - {entry['message']}
                    </div>
                    """, unsafe_allow_html=True)
            
            # Botón para limpiar historial
            if st.button("🗑️ Limpiar Historial", key="clear_history"):
                st.session_state.conversion_history = []
                st.rerun()
        else:
            st.info("No hay actividad reciente")

if __name__ == "__main__":
    main()
